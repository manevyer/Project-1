import os
import re
import requests
from bs4 import BeautifulSoup
import json
import logging
from urllib.parse import urljoin, urlparse
from collections import deque
from dotenv import load_dotenv
import io

# Optional imports for document parsing
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

try:
    from spire.doc import *
    from spire.doc.common import *
except ImportError:
    Document = None

# Langchain splitter
try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    RecursiveCharacterTextSplitter = None

# Set up logging for error handling and progress tracking
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')

def perform_login(session):
    """
    Attempts to log in to the METU system using credentials from .env via Basic Authentication.
    If no credentials are provided, simply returns False.
    """
    load_dotenv()
    username = os.getenv("METU_USERNAME")
    password = os.getenv("METU_PASSWORD")
    
    if not username or not password or username == "your_username":
        logging.info("No valid METU_USERNAME and METU_PASSWORD found in .env. Proceeding anonymously.")
        return False
        
    logging.info(f"Setting up Basic Authentication for {username}...")
    session.auth = (username, password)
    
    # Test authentication against a protected page
    test_url = "https://sp-ie.metu.edu.tr/en/forms"
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
    }
    
    try:
        test_response = session.get(test_url, headers=headers)
        if test_response.status_code == 200:
            logging.info("Login successful (Basic Authentication)!")
            return True
        elif test_response.status_code == 401:
            logging.error("Login failed! Incorrect username or password for Basic Auth.")
            return False
        else:
            logging.warning(f"Unexpected status code {test_response.status_code} during login check.")
            return False
            
    except Exception as e:
        logging.error(f"Login setup failed due to an error: {e}")
        return False

def fetch_html(url, session):
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
    }
    try:
        response = session.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        return BeautifulSoup(response.content, 'html.parser')
    except requests.exceptions.RequestException as e:
        logging.error(f"Failed to fetch HTML {url}: {e}")
        return None

def fetch_pdf_text(url, session):
    headers = {'User-Agent': 'Mozilla/5.0'}
    try:
        response = session.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        
        if not pypdf:
            logging.warning(f"pypdf not installed. Cannot read {url}")
            return ""
            
        pdf_file = io.BytesIO(response.content)
        reader = pypdf.PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
        return text
    except Exception as e:
        logging.error(f"Failed to read PDF {url}: {e}")
        return ""

def fetch_docx_text(url, session):
    headers = {'User-Agent': 'Mozilla/5.0'}
    try:
        response = session.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        
        if not docx:
            logging.warning(f"python-docx not installed. Cannot read {url}")
            return ""
            
        docx_file = io.BytesIO(response.content)
        doc = docx.Document(docx_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        logging.error(f"Failed to read DOCX {url}: {e}")
        return ""

def fetch_doc_text(url, session):
    headers = {'User-Agent': 'Mozilla/5.0'}
    try:
        response = session.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        
        if Document is None:
            logging.warning(f"Spire.Doc not installed. Cannot read {url}")
            return ""
            
        # Spire.Doc requires reading from a file object path on disk.
        temp_path = "temp_download.doc"
        with open(temp_path, "wb") as f:
            f.write(response.content)
            
        doc = Document()
        doc.LoadFromFile(temp_path)
        text = doc.GetText()
        doc.Close()
        
        if os.path.exists(temp_path):
            os.remove(temp_path)
            
        text = text.replace("Evaluation Warning: The document was created with Spire.Doc for Python.", "")
        return text.strip()
    except Exception as e:
        logging.error(f"Failed to read DOC {url}: {e}")
        return ""

def clean_data(soup):
    if not soup:
        return ""
    
    # Removed 'aside' to prevent data loss in important sidebars
    tags_to_remove = ['nav', 'footer', 'header', 'script', 'style', 'noscript', 'meta', 'link']
    for tag in tags_to_remove:
        for element in soup.find_all(tag):
            element.decompose()
            
    noise_classes = ['links', 'share-links', 'social', 'language-switcher', 'date-display-single', 'language', 'social-media']
    for cls in noise_classes:
        for element in soup.find_all(class_=cls):
            element.decompose()
            
    main_content = soup.find('main') or soup.find('article') or soup.find('div', class_='content') or soup.find('div', id='content') or soup.find('div', class_='region-content')
    
    if main_content:
        text = main_content.get_text(separator='\n', strip=True)
    else:
        text = soup.body.get_text(separator='\n', strip=True) if soup.body else soup.get_text(separator='\n', strip=True)
        
    # Improve whitespace cleaning to prevent formatting issues 
    text = re.sub(r'[\r\t\f\v ]+', ' ', text) # Replace horizontal spaces with a single space
    text = re.sub(r'\n\s*\n+', '\n\n', text)  # Replace multiple newlines with double newline to preserve paragraph separation
    
    return text.strip()

def chunk_text(text):
    if RecursiveCharacterTextSplitter:
        # Use LangChain Text Splitter to prevent breaking words and missing data optimally
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=700, 
            chunk_overlap=100,
            separators=["\n\n", "\n", " ", ""]
        )
        return splitter.split_text(text)
    else:
        # Fallback 
        raw_chunks = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for raw in raw_chunks:
            raw = raw.strip()
            if not raw:
                continue
            if len(current_chunk) + len(raw) < 500:
                if current_chunk:
                    current_chunk += "\n\n" + raw
                else:
                    current_chunk = raw
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                current_chunk = raw
                
        if current_chunk:
            chunks.append(current_chunk)
        return chunks

def generate_queries_with_ollama(chunk_text, max_retries=2):
    url = "http://localhost:11434/api/generate"
    prompt = f"""Based on the following text about university internships, write exactly 3 realistic questions that a student might ask to get this information.
Write 2 questions in Turkish and 1 question in English.
Provide ONLY the questions, one per line, without any numbering or extra text.

Text:
{chunk_text[:1000]}
"""
    payload = {
        "model": "llama3.2:latest",
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": 0.1
        }
    }
    
    for _ in range(max_retries):
        try:
            response = requests.post(url, json=payload, timeout=20)
            if response.status_code == 200:
                result = response.json().get('response', '')
                # Clean up the output to ensure we just get strings
                questions = [q.strip('- *."\'1234567890') for q in result.strip().split('\n') if q.strip()]
                if len(questions) >= 1:
                    return questions[:3]
        except Exception as e:
            logging.warning(f"Ollama query generation failed: {e}")
            break
            
    return None

def format_to_json(chunks, source_url, page_title):
    dataset = []
    
    for i, chunk in enumerate(chunks):
        topic = f"{page_title} - Section {i+1}" if page_title else f"Information from {source_url} - Section {i+1}"
        
        lower_chunk = chunk.lower()
        if 'ie 300' in lower_chunk or 'ie300' in lower_chunk:
            topic = "IE 300 Internship Requirements"
        elif 'ie 400' in lower_chunk or 'ie400' in lower_chunk:
            topic = "IE 400 Internship Requirements"
        elif 'document' in lower_chunk or 'form' in lower_chunk:
            topic = "Required Documents and Forms"
        elif 'deadline' in lower_chunk or 'date' in lower_chunk:
            topic = "Deadlines and Important Dates"
        elif 'step' in lower_chunk or 'apply' in lower_chunk or 'application' in lower_chunk:
            topic = "Application Steps and Procedures"
            
        logging.info(f"Generating queries with local LLM for chunk {i+1}/{len(chunks)}...")
        generated_questions = generate_queries_with_ollama(chunk)
        
        # Fallback to smart generic bilingual questions if Ollama fails or is not running
        if not generated_questions:
            generated_questions = [
                f"What is the information regarding {topic.lower()}?",
                f"Bana {topic} hakkında bilgi verebilir misiniz?",
                f"IE staj süreciyle ilgili {topic.lower()} detayları nelerdir?"
            ]

        entry = {
            "topic": topic,
            "user_query_variations": generated_questions,
            "chatbot_response": chunk
        }
        dataset.append(entry)
        
    return dataset

def get_all_internal_links(base_url, session):
    visited = set()
    queue = deque([base_url])
    all_links = set([base_url])
    base_domain = urlparse(base_url).netloc
    base_path = urlparse(base_url).path

    while queue:
        current_url = queue.popleft()
        if current_url in visited:
            continue
            
        visited.add(current_url)
        logging.info(f"Discovering links on: {current_url}")
        
        soup = fetch_html(current_url, session)
        if not soup:
            continue
            
        for a_tag in soup.find_all('a', href=True):
            href = a_tag['href']
            if href.startswith(('mailto:', 'tel:', 'javascript:', '#')):
                continue
                
            full_url = urljoin(current_url, href)
            full_url = full_url.split('#')[0]
            
            parsed_url = urlparse(full_url)
            
            if parsed_url.netloc == base_domain and parsed_url.path.startswith(base_path):
                # We skip obviously non-text extensions, but keep pdf and doc/docx
                if not full_url.lower().endswith(('.xls', '.xlsx', '.zip', '.rar', '.jpg', '.png', '.jpeg')):
                    if full_url not in all_links:
                        all_links.add(full_url)
                        
                        # Only append HTML pages to the queue, prevent fetching HTML tags from a PDF!
                        if not full_url.lower().endswith(('.pdf', '.doc', '.docx')):
                            queue.append(full_url)
                        
    return list(all_links)

def main():
    base_url = "https://sp-ie.metu.edu.tr/en"
    session = requests.Session()
    
    # 0. Attempt Login using environment variables
    perform_login(session)
    
    logging.info(f"Starting crawl to find all pages under: {base_url}")
    
    # 1. Discover all internal pages
    urls_to_scrape = get_all_internal_links(base_url, session)
    logging.info(f"Found {len(urls_to_scrape)} unique pages/documents to scrape.")
    
    final_dataset = []
    
    # 2. Iterate over ALL found pages and documents and extract data
    for url in urls_to_scrape:
        logging.info(f"Scraping content from: {url}")
        
        page_title = url.split('/')[-1] if not url.endswith('/') else url.split('/')[-2]
        cleaned_text = ""
        
        if url.lower().endswith('.pdf'):
            cleaned_text = fetch_pdf_text(url, session)
        elif url.lower().endswith('.docx'):
            cleaned_text = fetch_docx_text(url, session)
        elif url.lower().endswith('.doc'):
            cleaned_text = fetch_doc_text(url, session)
        else:
            soup = fetch_html(url, session)
            if soup:
                title_tag = soup.title
                page_title = title_tag.string.strip() if title_tag and title_tag.string else page_title
                cleaned_text = clean_data(soup)
        
        if cleaned_text:
            chunks = chunk_text(cleaned_text)
            json_data = format_to_json(chunks, url, page_title)
            final_dataset.extend(json_data)
        else:
            logging.warning(f"No content extracted from {url}")

    # 3. Save to JSON file locally
    output_filename = "metu_ie_chatbot_dataset.json"
    try:
        with open(output_filename, 'w', encoding='utf-8') as f:
            json.dump(final_dataset, f, ensure_ascii=False, indent=2)
        logging.info(f"Successfully scraped all information and saved {len(final_dataset)} entries to {output_filename}")
    except IOError as e:
        logging.error(f"Failed to save JSON file: {e}")

if __name__ == "__main__":
    main()